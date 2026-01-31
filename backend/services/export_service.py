"""
Nexus AI - Export Service
Handles project data conversion to various formats: Markdown, JSON, PDF, DOCX
"""

import json
import io
import os
from datetime import datetime
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches

class ExportService:
    @staticmethod
    def to_markdown(project_data):
        """Convert project data to Markdown."""
        md = f"# Project: {project_data['name']}\n\n"
        md += f"**Description**: {project_data.get('description', 'No description')}\n\n"
        md += f"**Status**: {project_data['status']}\n"
        md += f"**Progress**: {project_data['progress']}%\n"
        md += f"**Total Phases**: {project_data['total_phases']}\n"
        md += f"**Total Tasks**: {project_data['total_tasks']}\n"
        md += f"**Estimated Minutes**: {project_data['estimated_minutes']}\n"
        md += f"**Risk Level**: {project_data['risk_level']}\n\n"

        if project_data.get('project_plan'):
            md += "## Project Plan\n\n"
            for phase in project_data['project_plan']:
                md += f"### Phase {phase.get('phase_number', '?')}: {phase.get('phase_name', 'Unnamed')}\n"
                for task in phase.get('tasks', []):
                    md += f"- **Task {task.get('task_id', '?')}**: {task.get('description', 'No description')} ({task.get('status', 'pending')})\n"
                md += "\n"

        if project_data.get('output'):
            md += "## Execution Output\n\n"
            md += project_data['output'] + "\n\n"

        if project_data.get('summary'):
            md += "## Summary\n\n"
            md += project_data['summary'] + "\n"

        return md

    @staticmethod
    def to_json(project_data):
        """Convert project data to JSON."""
        return json.dumps(project_data, indent=4, default=str)

    @staticmethod
    def to_pdf(project_data):
        """Generate PDF for project data."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER)
        styles = getSampleStyleSheet()
        story = []

        # Title
        story.append(Paragraph(f"Project: {project_data['name']}", styles['Title']))
        story.append(Spacer(1, 12))

        # Metadata
        story.append(Paragraph(f"<b>Description</b>: {project_data.get('description', 'No description')}", styles['Normal']))
        story.append(Paragraph(f"<b>Status</b>: {project_data['status']}", styles['Normal']))
        story.append(Paragraph(f"<b>Progress</b>: {project_data['progress']}%", styles['Normal']))
        story.append(Paragraph(f"<b>Estimated Time</b>: {project_data['estimated_minutes']} minutes", styles['Normal']))
        story.append(Spacer(1, 12))

        # Plan
        if project_data.get('project_plan'):
            story.append(Paragraph("Project Plan", styles['Heading2']))
            data = [["Phase", "Task ID", "Description", "Status"]]
            for phase in project_data['project_plan']:
                for task in phase.get('tasks', []):
                    data.append([
                        phase.get('phase_name', ''),
                        task.get('task_id', ''),
                        task.get('description', ''),
                        task.get('status', '')
                    ])
            
            if len(data) > 1:
                t = Table(data, colWidths=[1.5*inch, 1*inch, 3.5*inch, 1*inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 24))

        # Output
        if project_data.get('output'):
            story.append(Paragraph("Execution Output", styles['Heading2']))
            # Simple text fix for PDF
            clean_output = str(project_data['output']).replace('\n', '<br/>')
            story.append(Paragraph(clean_output, styles['Normal']))
            story.append(Spacer(1, 12))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    @staticmethod
    def to_docx(project_data):
        """Generate DOCX for project data."""
        doc = Document()
        doc.add_heading(f"Project: {project_data['name']}", 0)

        p = doc.add_paragraph()
        p.add_run('Description: ').bold = True
        p.add_run(project_data.get('description', 'No description'))

        doc.add_paragraph(f"Status: {project_data['status']}")
        doc.add_paragraph(f"Progress: {project_data['progress']}%")

        if project_data.get('project_plan'):
            doc.add_heading('Project Plan', level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Phase'
            hdr_cells[1].text = 'Task ID'
            hdr_cells[2].text = 'Description'
            hdr_cells[3].text = 'Status'

            for phase in project_data['project_plan']:
                for task in phase.get('tasks', []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(phase.get('phase_name', ''))
                    row_cells[1].text = str(task.get('task_id', ''))
                    row_cells[2].text = str(task.get('description', ''))
                    row_cells[3].text = str(task.get('status', ''))

        if project_data.get('output'):
            doc.add_heading('Execution Output', level=1)
            doc.add_paragraph(project_data['output'])

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
